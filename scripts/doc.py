import argparse
import sys
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def clean_text(text):
    """
    智能清理文本：
    1. 移除Markdown内联格式标记
    2. 智能处理空格：
       - 保留英文单词间的半角空格（两侧都是ASCII字符）
       - 删除中文文本中的空格（包括中英文之间的空格）
    3. 删除全角空格（中文空格）
    4. 删除所有制表符
    5. 合并多个连续空格为单个空格
    6. 移除首尾空白
    7. 从源头减少隐藏空格问题
    """
    # 移除Markdown内联格式标记
    # 移除粗体标记：**文本** 或 __文本__
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # 移除斜体标记：*文本* 或 _文本_
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # 移除行内代码标记：`代码`
    text = re.sub(r'`(.+?)`', r'\1', text)
    
    # 删除所有全角空格（中文空格）
    text = text.replace('　', '')
    
    # 删除所有制表符
    text = text.replace('\t', ' ')
    
    # 智能空格处理：保留英文上下文中的空格，删除中文中的空格
    # 我们将文本分割成单词和空格，然后决定哪些空格保留
    # 简单方法：遍历字符，构建新文本
    new_text_chars = []
    i = 0
    while i < len(text):
        char = text[i]
        if char == ' ':
            # 检查空格前后字符
            prev_char = text[i-1] if i > 0 else ''
            next_char = text[i+1] if i < len(text)-1 else ''
            
            # 判断是否保留空格：前后都是ASCII可打印字符（英文上下文）
            # ASCII可打印字符：32-126（包括空格、标点、字母、数字）
            prev_is_ascii = prev_char and 32 <= ord(prev_char) <= 126
            next_is_ascii = next_char and 32 <= ord(next_char) <= 126
            
            # 如果空格在英文上下文之间，保留它
            if prev_is_ascii and next_is_ascii:
                new_text_chars.append(' ')
            # 否则删除空格（中文中的空格或中英文之间的空格）
            i += 1
            continue
        else:
            new_text_chars.append(char)
            i += 1
    
    text = ''.join(new_text_chars)
    
    # 合并多个连续空格为单个空格（如果还有）
    text = re.sub(r'[ ]{2,}', ' ', text)
    
    # 移除首尾空白（包括空格）
    text = text.strip()
    
    return text

def set_run_font_completely(run, font_name='宋体', size_pt=12, bold=False, italic=False):
    """
    完全设置运行的字体属性，确保所有字符集都使用指定字体
    解决中英文、符号字体不一致问题
    """
    # 基本字体设置
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 尝试设置字符集特定的字体（解决中英文字体不一致问题）
    try:
        # 获取运行的XML元素
        r = run._r
        rPr = r.get_or_add_rPr()
        
        # 创建字体设置元素
        rFonts = OxmlElement('w:rFonts')
        
        # 设置所有字符集为宋体
        # ascii: 西文字符
        # eastAsia: 东亚字符
        # hAnsi: 西文字符（另一种）
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)  # 复杂脚本
        
        # 将字体设置添加到运行属性中
        rPr.append(rFonts)
    except Exception as e:
        # 如果XML操作失败，继续使用常规设置
        pass

def parse_markdown_to_docx(md_content, doc):
    """
    将 Markdown 文本转换为 Word 文档内容（改进版）。
    支持：
        - # 标题1 到 ###### 标题6（支持前导空格和尾部#号）
        - 无序列表（- 或 *）
        - 有序列表（数字. ）
        - 普通段落
        - 移除内联格式标记（**粗体**、*斜体*）
    
    重要改进：确保每个段落只有一个文本运行，避免隐藏字符导致字体不一致。
    """
    # 匹配 ATX 风格标题：可选空格 + 1-6个# + 空格 + 标题内容 + 可选尾部#号
    atx_heading_re = re.compile(r'^(\s*)(#{1,6})\s+(.+?)\s*#*\s*$')
    

    
    # 函数：为段落设置统一的字体
    def set_paragraph_font(paragraph, is_heading=False, level=1):
        """确保段落及其所有运行使用宋体"""
        # 首先设置段落样式字体（如果有）
        if paragraph.style and paragraph.style.font:
            paragraph.style.font.name = '宋体'
            paragraph.style.font.italic = False
            if not is_heading:
                paragraph.style.font.bold = False
        
        # 确保所有运行使用宋体
        for run in paragraph.runs:
            run.font.name = '宋体'
            run.font.italic = False
            # 如果是标题，保持粗体；否则取消粗体
            if is_heading:
                run.font.bold = True
            else:
                run.font.bold = False
            # 确保颜色为黑色
            if run.font.color:
                run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置字体大小
            if not run.font.size:
                if is_heading:
                    if level == 1:
                        run.font.size = Pt(16)
                    elif level == 2:
                        run.font.size = Pt(14)
                    elif level == 3:
                        run.font.size = Pt(13)
                    else:
                        run.font.size = Pt(12)
                else:
                    run.font.size = Pt(12)
    
    # 函数：设置段落格式（减少段落间距）
    def set_paragraph_format(paragraph, is_heading=False):
        """设置段落格式，减少段落间距"""
        format = paragraph.paragraph_format
        
        # 设置段落间距
        if is_heading:
            # 标题：段前6磅，段后3磅
            format.space_before = Pt(6)
            format.space_after = Pt(3)
        else:
            # 正文：段前0磅，段后0磅，行距1.5倍
            format.space_before = Pt(0)
            format.space_after = Pt(0)
            format.line_spacing = 1.5
        
        # 设置首行缩进（正文首行缩进2字符，约8磅）
        if not is_heading and not paragraph.style.name.startswith('List'):
            # 正文段落首行缩进
            format.first_line_indent = Pt(16)  # 约2字符
    
    # 函数：创建单运行段落
    def create_single_run_paragraph(doc, text, style=None, is_heading=False, level=1):
        """创建一个只有一个文本运行的段落，确保字体统一"""
        if style:
            paragraph = doc.add_paragraph(style=style)
        elif is_heading:
            paragraph = doc.add_heading(level=min(level, 6))
        else:
            paragraph = doc.add_paragraph()
        
        # 添加文本运行
        run = paragraph.add_run(text)
        
        # 使用完全字体设置函数，确保所有字符集都使用宋体
        if is_heading:
            # 标题字体大小根据级别调整
            if level == 1:
                size_pt = 16
            elif level == 2:
                size_pt = 14
            elif level == 3:
                size_pt = 13
            else:
                size_pt = 12
            set_run_font_completely(run, font_name='宋体', size_pt=size_pt, bold=True, italic=False)
        else:
            # 正文：小四（12磅），无粗体
            set_run_font_completely(run, font_name='宋体', size_pt=12, bold=False, italic=False)
        
        # 同时设置段落样式（保持向后兼容）
        set_paragraph_font(paragraph, is_heading, level)
        
        # 设置段落格式（减少间距）
        set_paragraph_format(paragraph, is_heading)
        
        return paragraph
    
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        
        # 尝试匹配 ATX 标题
        match = atx_heading_re.match(line)
        if match:
            # 计算标题级别（#号数量）
            level = len(match.group(2))
            # 获取标题文本，去除两端空白
            text = match.group(3).strip()
            # 清理文本：移除内联格式标记和多余空格
            text = clean_text(text)
            # 使用单运行创建标题
            p = create_single_run_paragraph(doc, text, is_heading=True, level=level)
            i += 1
            continue
        
        # 无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            text = clean_text(text)
            p = create_single_run_paragraph(doc, text, style='List Bullet')
        # 有序列表（简单处理，以数字开头后跟点号）
        elif re.match(r'^\d+\.\s+', stripped):
            # 去除数字和点号
            text = re.sub(r'^\d+\.\s+', '', stripped)
            text = clean_text(text)
            p = create_single_run_paragraph(doc, text, style='List Number')
        else:
            # 普通段落，清理文本
            text = clean_text(stripped)
            p = create_single_run_paragraph(doc, text)
        i += 1

def main():
    parser = argparse.ArgumentParser(description='从 Markdown 生成 Word 文档')
    parser.add_argument('--title', required=True, help='文档标题')
    parser.add_argument('--input', required=True, help='输入的 Markdown 文件路径')
    parser.add_argument('--output', required=True, help='输出的 Word 文档路径')
    args = parser.parse_args()

    # 检查输入文件是否存在
    if not os.path.exists(args.input):
        print(f"错误：输入文件 {args.input} 不存在", file=sys.stderr)
        sys.exit(1)

    # 读取 Markdown 内容
    with open(args.input, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 创建 Word 文档
    doc = Document()

    # 设置文档默认样式：宋体、小四（12磅）、黑色、无斜体
    # 1. 设置默认段落字体（影响所有未明确设置的文本）
    if 'Default Paragraph Font' in doc.styles:
        default_font = doc.styles['Default Paragraph Font'].font
        default_font.name = '宋体'
        default_font.size = Pt(12)
    
    # 2. 设置正文样式（Normal）
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '宋体'
    normal_font.size = Pt(12)  # 小四对应12磅
    normal_font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    normal_font.italic = False  # 确保无斜体
    normal_font.bold = False    # 确保无粗体（除非需要）
    
    # 3. 设置所有标题样式（Heading 1 到 Heading 9）
    for level in range(1, 10):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_font = heading_style.font
            heading_font.name = '宋体'
            heading_font.color.rgb = RGBColor(0, 0, 0)
            heading_font.italic = False  # 清除斜体
            # 标题字号根据级别调整（可自定义）
            if level == 1:
                heading_font.size = Pt(16)  # 一号字
            elif level == 2:
                heading_font.size = Pt(14)  # 小二
            elif level == 3:
                heading_font.size = Pt(13)  # 小三
            else:
                heading_font.size = Pt(12)  # 小四
    
    # 4. 设置其他重要样式
    important_styles = ['Title', 'Subtitle', 'Quote', 'Intense Quote', 'Caption', 
                       'Body Text', 'Body Text 2', 'Body Text 3', 'macro',
                       'Strong', 'Emphasis', 'Book Title', 'TOC Heading',
                       'Footnote Text', 'Header', 'Footer', 'Page Number']
    
    for style_name in important_styles:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            font = style.font
            font.name = '宋体'  # 设置为宋体
            font.italic = False  # 清除斜体
            # 保持原有字号，或设置默认字号
            if not font.size:
                font.size = Pt(12)
    
    # 5. 设置列表样式
    for style_name in ['List Bullet', 'List Number']:
        if style_name in doc.styles:
            list_style = doc.styles[style_name]
            list_font = list_style.font
            list_font.name = '宋体'
            list_font.size = Pt(12)  # 小四
            list_font.color.rgb = RGBColor(0, 0, 0)
            list_font.italic = False  # 清除斜体

    # 添加标题 - 使用单运行方法确保字体统一
    # 创建Title样式的段落
    title_para = doc.add_paragraph(style='Title')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(clean_text(args.title))
    # 使用完全字体设置函数，确保标题所有字符都使用宋体
    set_run_font_completely(title_run, font_name='宋体', size_pt=18, bold=True, italic=False)

    # 解析并添加内容
    parse_markdown_to_docx(md_content, doc)
    
    # 6. 额外步骤：遍历所有段落和运行，确保字体设置完全生效
    # 使用完全字体设置函数覆盖所有可能的字体设置
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # 根据段落样式确定字体参数
            style_name = paragraph.style.name if paragraph.style else ''
            
            # 判断是否为标题
            is_heading = style_name.startswith('Heading') or style_name == 'Title'
            
            # 判断标题级别（如果适用）
            level = 1
            if style_name.startswith('Heading'):
                try:
                    # 提取标题级别，如"Heading 1" -> 1
                    level_match = re.search(r'Heading\s+(\d+)', style_name)
                    if level_match:
                        level = int(level_match.group(1))
                except:
                    level = 1
            
            # 设置字体大小
            if is_heading:
                if level == 1:
                    size_pt = 16
                elif level == 2:
                    size_pt = 14
                elif level == 3:
                    size_pt = 13
                elif style_name == 'Title':
                    size_pt = 18
                else:
                    size_pt = 12
            else:
                size_pt = 12  # 小四
            
            # 使用完全字体设置函数
            set_run_font_completely(
                run, 
                font_name='宋体', 
                size_pt=size_pt, 
                bold=is_heading,  # 标题加粗
                italic=False
            )
    
    # 7. 设置文档核心默认字体（尝试设置西文字体也为宋体）
    # 在Word中，可以通过设置样式的中西文字体来解决
    # 这里我们尝试设置所有可能的默认字体
    
    # 尝试设置文档的默认字符集
    try:
        # 获取文档的样式部分
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # 获取文档的样式元素
        style_element = doc.styles.element
        
        # 尝试设置默认字体
        # 这种方法更底层，但可能更有效
        pass
    except:
        # 如果失败，继续使用常规方法
        pass

    # 保存文档
    doc.save(args.output)
    print(f"文档已生成：{args.output}")

if __name__ == '__main__':
    main()